import os
import base64
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()


class FileAgent:
    """
    Reads and analyses ALL file types:
    CSV, Excel, Word, PDF, Images, Text
    """

    def __init__(self):
        self.hf_token = os.getenv(
            "HUGGINGFACE_TOKEN", ""
        )

    def detect_type(self, path: str) -> str:
        ext = Path(path).suffix.lower()
        types = {
            ".csv":  "csv",
            ".xlsx": "excel",
            ".xls":  "excel",
            ".docx": "word",
            ".doc":  "word",
            ".pdf":  "pdf",
            ".txt":  "text",
            ".md":   "text",
            ".jpg":  "image",
            ".jpeg": "image",
            ".png":  "image",
            ".gif":  "image",
            ".bmp":  "image",
            ".webp": "image",
        }
        return types.get(ext, "unknown")

    def analyse(
        self,
        path: str,
        question: str = None
    ) -> dict:
        file_type = self.detect_type(path)
        file_name = Path(path).name
        file_size = round(
            os.path.getsize(path) / 1024, 2
        )

        if file_type == "csv":
            result = self._read_csv(path)
        elif file_type == "excel":
            result = self._read_excel(path)
        elif file_type == "word":
            result = self._read_word(path)
        elif file_type == "pdf":
            result = self._read_pdf(path)
        elif file_type == "image":
            result = self._read_image(
                path, question
            )
        elif file_type == "text":
            result = self._read_text(path)
        else:
            result = {
                "status": "error",
                "error": (
                    f"File type not supported. "
                    f"Supported: CSV, Excel, "
                    f"Word, PDF, Images, Text"
                )
            }

        result["file_name"]    = file_name
        result["file_size_kb"] = file_size
        return result

    def _read_csv(self, path: str) -> dict:
        try:
            import pandas as pd
            df = pd.read_csv(path)

            numeric_cols = df.select_dtypes(
                include="number"
            ).columns.tolist()

            text_cols = df.select_dtypes(
                include="object"
            ).columns.tolist()

            stats = {}
            for col in numeric_cols[:5]:
                stats[col] = {
                    "mean": round(
                        float(df[col].mean()), 2
                    ),
                    "min":  round(
                        float(df[col].min()), 2
                    ),
                    "max":  round(
                        float(df[col].max()), 2
                    )
                }

            return {
                "status":       "success",
                "type":         "csv",
                "rows":         len(df),
                "columns":      len(df.columns),
                "column_names": df.columns.tolist(),
                "numeric_cols": numeric_cols,
                "text_cols":    text_cols,
                "missing":      int(
                    df.isnull().sum().sum()
                ),
                "statistics":   stats,
                "preview":      df.head(3).to_dict(),
                "summary": (
                    f"CSV file with {len(df)} rows "
                    f"and {len(df.columns)} columns. "
                    f"Numeric columns: "
                    f"{', '.join(numeric_cols[:3])}. "
                    f"Text columns: "
                    f"{', '.join(text_cols[:3])}. "
                    f"Missing values: "
                    f"{df.isnull().sum().sum()}"
                )
            }
        except Exception as e:
            return {
                "status": "error",
                "type":   "csv",
                "error":  str(e)
            }

    def _read_excel(self, path: str) -> dict:
        try:
            import pandas as pd
            xl     = pd.ExcelFile(path)
            sheets = {}

            for sheet in xl.sheet_names[:5]:
                df = pd.read_excel(
                    path, sheet_name=sheet
                )
                sheets[sheet] = {
                    "rows":         len(df),
                    "columns":      len(df.columns),
                    "column_names": (
                        df.columns.tolist()
                    ),
                    "missing":      int(
                        df.isnull().sum().sum()
                    ),
                    "numeric_cols": (
                        df.select_dtypes(
                            include="number"
                        ).columns.tolist()
                    )
                }

            total_rows = sum(
                s["rows"] for s in sheets.values()
            )

            return {
                "status":       "success",
                "type":         "excel",
                "total_sheets": len(xl.sheet_names),
                "sheet_names":  xl.sheet_names,
                "sheets":       sheets,
                "summary": (
                    f"Excel file with "
                    f"{len(xl.sheet_names)} sheet(s) "
                    f"and {total_rows} total rows."
                )
            }
        except Exception as e:
            return {
                "status": "error",
                "type":   "excel",
                "error":  str(e)
            }

    def _read_word(self, path: str) -> dict:
        try:
            from docx import Document
            doc = Document(path)

            paragraphs = [
                p.text for p in doc.paragraphs
                if p.text.strip()
            ]
            full_text = "\n".join(paragraphs)

            tables_data = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    row_text = " | ".join([
                        c.text.strip()
                        for c in row.cells
                    ])
                    if row_text.strip():
                        rows.append(row_text)
                if rows:
                    tables_data.append(rows[:5])

            summary = self._summarize(
                full_text[:2000]
            )

            return {
                "status":      "success",
                "type":        "word",
                "summary":     summary,
                "full_text":   full_text[:4000],
                "paragraphs":  len(paragraphs),
                "tables":      len(doc.tables),
                "tables_data": tables_data[:2],
                "word_count":  len(
                    full_text.split()
                ),
                "char_count":  len(full_text)
            }
        except Exception as e:
            return {
                "status": "error",
                "type":   "word",
                "error":  str(e)
            }

    def _read_pdf(self, path: str) -> dict:
        try:
            import fitz
            import re
            from collections import Counter

            doc          = fitz.open(path)
            pages        = len(doc)
            full_text    = ""
            images_found = 0
            tables_found = 0
            page_texts   = []

            for i, page in enumerate(doc):

                # Extract text from each page
                page_text = page.get_text()
                full_text += (
                    f"\n[Page {i+1}]\n"
                    + page_text
                )
                page_texts.append({
                    "page":  i + 1,
                    "chars": len(page_text),
                    "words": len(page_text.split())
                })

                # Count images on each page
                images_found += len(
                    page.get_images()
                )

                # Detect tables by looking for
                # grid like patterns in text
                lines = page_text.split("\n")
                table_lines = [
                    l for l in lines
                    if len(l.split()) > 3
                    and any(
                        c in l
                        for c in ["|", "\t", "  "]
                    )
                ]
                if len(table_lines) > 3:
                    tables_found += 1

            doc.close()

            # Extract key topics from the text
            words = re.findall(
                r'\b[A-Z][a-z]+\b',
                full_text[:3000]
            )
            common_words = [
                w for w, c
                in Counter(words).most_common(10)
                if len(w) > 4
            ]

            # Generate AI summary using TinyLlama
            summary = self._summarize(
                full_text[:3000]
            )

            # Generate detailed analysis
            detailed_analysis = (
                self._analyse_pdf_content(
                    full_text[:5000]
                )
            )

            return {
                "status":             "success",
                "type":               "pdf",
                "summary":            summary,
                "detailed_analysis":  detailed_analysis,
                "full_text":          full_text[:5000],
                "pages":              pages,
                "images":             images_found,
                "tables":             tables_found,
                "word_count":         len(
                    full_text.split()
                ),
                "char_count":         len(full_text),
                "page_details":       page_texts,
                "key_topics":         common_words,
                "avg_words_per_page": round(
                    len(full_text.split()) / pages
                ) if pages > 0 else 0
            }

        except Exception as e:
            return {
                "status": "error",
                "type":   "pdf",
                "error":  str(e)
            }

    def _analyse_pdf_content(
        self,
        text: str
    ) -> str:
        """Generate detailed analysis of PDF content"""
        try:
            import ollama
            response = ollama.chat(
                model="tinyllama",
                messages=[{
                    "role": "user",
                    "content": (
                        f"Analyse this document "
                        f"content and provide:\n"
                        f"1. Main topic of the "
                        f"document\n"
                        f"2. Key points covered\n"
                        f"3. Type of document "
                        f"(report, article, "
                        f"manual etc)\n"
                        f"4. Any important numbers "
                        f"or statistics found\n\n"
                        f"Document content:\n{text}"
                    )
                }]
            )
            return response["message"]["content"]
        except Exception:
            return (
                "Detailed analysis not available. "
                "Basic summary has been generated."
            )

    def _read_image(
        self,
        path: str,
        question: str = None
    ) -> dict:
        try:
            from PIL import Image as PILImage

            img    = PILImage.open(path)
            width  = img.width
            height = img.height
            mode   = img.mode
            size   = os.path.getsize(path)

            if question is None:
                question = (
                    "Describe this image in detail. "
                    "Include: what you see, "
                    "all objects present, "
                    "any text visible, "
                    "colors and overall context."
                )

            description = self._analyse_image_hf(
                path, question
            )

            return {
                "status":      "success",
                "type":        "image",
                "description": description,
                "metadata": {
                    "width":   width,
                    "height":  height,
                    "mode":    mode,
                    "size_kb": round(
                        size / 1024, 2
                    ),
                    "format":  Path(
                        path
                    ).suffix.upper()
                }
            }
        except Exception as e:
            return {
                "status": "error",
                "type":   "image",
                "error":  str(e)
            }

    def _analyse_image_hf(
        self,
        path: str,
        question: str
    ) -> str:
        try:
            from PIL import Image as PILImage

            img    = PILImage.open(path)
            width  = img.width
            height = img.height
            mode   = img.mode
            fmt    = Path(path).suffix.upper()
            size   = round(
                os.path.getsize(path) / 1024, 2
            )

            # Basic image info without AI
            description = (
                f"Image successfully uploaded "
                f"and processed.\n\n"
                f"Image details:\n"
                f"- Format: {fmt}\n"
                f"- Size: {size} KB\n"
                f"- Dimensions: "
                f"{width} x {height} px\n"
                f"- Color mode: {mode}\n\n"
                f"Note: Advanced AI image analysis "
                f"requires a paid HuggingFace "
                f"account. Basic metadata has "
                f"been extracted successfully."
            )

            return description

        except Exception as e:
            return (
                f"Image processing failed: {str(e)}"
            )

    def _read_text(self, path: str) -> dict:
        try:
            with open(
                path, "r",
                encoding="utf-8",
                errors="ignore"
            ) as f:
                content = f.read()

            summary = self._summarize(
                content[:2000]
            )

            return {
                "status":     "success",
                "type":       "text",
                "summary":    summary,
                "full_text":  content[:5000],
                "word_count": len(
                    content.split()
                ),
                "line_count": len(
                    content.splitlines()
                ),
                "char_count": len(content)
            }
        except Exception as e:
            return {
                "status": "error",
                "type":   "text",
                "error":  str(e)
            }

    def _summarize(self, text: str) -> str:
        try:
            import ollama
            response = ollama.chat(
                model="tinyllama",
                messages=[{
                    "role": "user",
                    "content": (
                        f"Summarize this content "
                        f"in 3 clear sentences. "
                        f"Be concise and informative:"
                        f"\n\n{text}"
                    )
                }]
            )
            return response["message"]["content"]
        except Exception:
            words = text.split()[:100]
            return (
                " ".join(words) + "..."
                if len(text.split()) > 100
                else text
            )